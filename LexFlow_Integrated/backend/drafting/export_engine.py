from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from xhtml2pdf import pisa
import markdown
import re

def export_to_word(markdown_text: str):
    """
    Converts Markdown to a clean, professional Word Doc.
    """
    doc = Document()
    
    # 1. Set Default Legal Font (Times New Roman)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    
    # Title
    title = doc.add_heading('LEGAL DRAFT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.name = 'Times New Roman'
    title.style.font.color.rgb = RGBColor(0, 0, 0)
    title.style.paragraph_format.space_after = Pt(12)

    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line: continue
            
        # Headers
        if line.startswith('### '):
            h = doc.add_heading(line.replace('### ', ''), level=2)
            h.style.font.color.rgb = RGBColor(0, 0, 0)
            h.style.font.name = 'Times New Roman'
            h.style.font.size = Pt(13)
            h.style.font.bold = True
        elif line.startswith('## '):
            h = doc.add_heading(line.replace('## ', ''), level=1)
            h.style.font.color.rgb = RGBColor(0, 0, 0)
            h.style.font.name = 'Times New Roman'
            h.style.font.size = Pt(14)
            h.style.font.bold = True

        # Citations
        elif line.startswith('> '):
            p = doc.add_paragraph(line.replace('> ', ''))
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].italic = True
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')

        # Normal Text
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            parts = line.split('**')
            for i, part in enumerate(parts):
                runner = p.add_run(part)
                runner.font.name = 'Times New Roman'
                if i % 2 == 1: runner.bold = True

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_to_pdf(markdown_text: str):
    """
    Converts Markdown -> HTML -> PDF.
    Fixed CSS to prevent 'NotImplementedType' crash.
    """
    html_body = markdown.markdown(markdown_text)
    
    result = BytesIO()
    
    # CSS: Simplified @page to prevent crash
    styled_html = f"""
    <html>
    <head>
        <style>
            @page {{
                size: A4;
                margin: 2.5cm;
            }}
            body {{ 
                font-family: Helvetica, Arial, sans-serif;
                font-size: 11pt; 
                line-height: 1.3; 
                color: #000000;
                text-align: justify;
            }}
            h1, h2, h3 {{ 
                color: #000000; 
                margin-top: 15px; 
                margin-bottom: 5px; 
                text-transform: uppercase;
                font-size: 12pt;
                font-weight: bold;
                text-decoration: underline;
            }}
            p {{ 
                margin-top: 0px;
                margin-bottom: 8px; 
            }}
            ul {{ margin-bottom: 8px; padding-left: 20px; }}
            li {{ margin-bottom: 2px; }}
            blockquote {{ 
                font-style: italic; 
                margin-left: 30px; 
                margin-top: 5px;
                margin-bottom: 5px;
                color: #333;
                border-left: 2px solid #ccc;
                padding-left: 10px;
            }}
        </style>
    </head>
    <body>
        <div style="text-align:center; font-weight:bold; font-size:14pt; margin-bottom:20px;">
            LEGAL DRAFT
        </div>
        <hr style="border-top: 1px solid #000; margin-bottom: 20px;">
        {html_body}
    </body>
    </html>
    """
    
    pisa.CreatePDF(styled_html, dest=result)
    result.seek(0)
    return result